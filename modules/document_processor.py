import os
import re
import fitz  # PyMuPDF
import pytesseract

from PIL import Image
from docx import Document


# Tell pytesseract where Tesseract is installed
pytesseract.pytesseract.tesseract_cmd = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
)


def clean_text(text):
    """Clean extracted text."""

    # Normalize line breaks
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Remove excessive spaces/tabs
    text = re.sub(r"[ \t]+", " ", text)

    # Remove excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def extract_from_pdf(file_path):
    """
    Extract text from a PDF.

    First tries normal PDF text extraction.
    If a page has no selectable text, OCR is used.
    """

    text = ""

    pdf = fitz.open(file_path)

    for page_number, page in enumerate(pdf, start=1):

        # -----------------------------------------
        # STEP 1: Normal PDF text extraction
        # -----------------------------------------
        page_text = page.get_text("text")

        if page_text.strip():

            text += f"\n[Page {page_number}]\n"
            text += page_text

        else:

            # -----------------------------------------
            # STEP 2: OCR fallback
            # -----------------------------------------

            print(
                f"OCR required for page {page_number}..."
            )

            # Render PDF page as an image
            pix = page.get_pixmap(
                matrix=fitz.Matrix(2, 2),
                alpha=False
            )

            # Convert pixmap to PIL image
            image = Image.frombytes(
                "RGB",
                [pix.width, pix.height],
                pix.samples
            )

            # Run OCR
            ocr_text = pytesseract.image_to_string(
                image
            )

            if ocr_text.strip():

                text += f"\n[Page {page_number} - OCR]\n"
                text += ocr_text

    pdf.close()

    return clean_text(text)


def extract_from_docx(file_path):
    """Extract text from a DOCX file."""

    document = Document(file_path)

    text_parts = []

    # Extract paragraphs
    for paragraph in document.paragraphs:

        if paragraph.text.strip():

            text_parts.append(
                paragraph.text
            )

    # Extract tables
    for table in document.tables:

        for row in table.rows:

            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
            )

            if row_text.strip():

                text_parts.append(
                    row_text
                )

    return clean_text(
        "\n".join(text_parts)
    )


def extract_from_txt(file_path):
    """Extract text from a TXT file."""

    with open(
        file_path,
        "r",
        encoding="utf-8",
        errors="ignore"
    ) as file:

        text = file.read()

    return clean_text(text)


def extract_text(file_path):
    """
    Extract text based on file type.

    Supported:
    PDF
    DOCX
    TXT
    """

    extension = os.path.splitext(
        file_path
    )[1].lower()

    if extension == ".pdf":

        return extract_from_pdf(
            file_path
        )

    elif extension == ".docx":

        return extract_from_docx(
            file_path
        )

    elif extension == ".txt":

        return extract_from_txt(
            file_path
        )

    else:

        raise ValueError(
            "Unsupported file type. "
            "Please upload a PDF, DOCX, or TXT file."
        )


if __name__ == "__main__":

    file_path = "test_report.pdf"

    try:

        text = extract_text(
            file_path
        )

        print(
            "\n===== EXTRACTED TEXT =====\n"
        )

        print(
            text[:5000]
        )

        print(
            "\n===== EXTRACTION SUCCESSFUL ====="
        )

    except Exception as e:

        print(
            f"\nERROR: {e}"
        )